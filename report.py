"""
report.py — reads performance xlsx + CDD sync xlsx → Word .docx + Slack text
Claude writes the conclusion paragraph and the Slack summary.
"""
import logging
import os
from collections import defaultdict
from datetime import timedelta, datetime

import anthropic
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

log = logging.getLogger(__name__)

# ── style constants ────────────────────────────────────────────────────────────
HDR_FILL   = "1A6496"   # steel blue  — matches tg_smc reference
ALT_FILL   = "EBF3FB"   # pale blue   — matches tg_smc reference
TITLE_RGB  = RGBColor(0x17, 0x36, 0x5D)   # dark navy
GREY_RGB   = RGBColor(0x66, 0x66, 0x66)
FONT       = "Times New Roman"

# Coverage bands: HIGH >=95% (green) | MODERATE 70-95% (orange) | LOW <70% (red)
STATUS_COLOR = {
    "HIGH":         RGBColor(0x1A, 0x7A, 0x1A),   # green
    "MODERATE":     RGBColor(0xE0, 0x60, 0x00),   # orange
    "LOW":          RGBColor(0xCC, 0x00, 0x00),   # red
    "NO TARGET":    RGBColor(0x88, 0x88, 0x88),
    "LOW ACTIVITY": RGBColor(0x88, 0x88, 0x88),
}

# Cell background fills for coverage colour-coding
COV_FILL = {
    "HIGH":         "C6EFCE",   # light green
    "MODERATE":     "FFEB9C",   # light amber
    "LOW":          "FFC7CE",   # light red
    "NO TARGET":    "F2F2F2",
    "LOW ACTIVITY": "F2F2F2",
}

def _cov_band(pct):
    """Return band name for a coverage percentage."""
    if pct >= 95:  return "HIGH"
    if pct >= 70:  return "MODERATE"
    return "LOW"


# ── docx helpers (verbatim pattern from generate_day4_report.py) ───────────────

def _add_hyperlink(para, text, url):
    """Insert a clickable hyperlink run into an existing paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    para._p.append(hyperlink)


def set_cell_bg(cell, hex6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    ex = tcPr.find(qn("w:shd"))
    if ex is not None:
        tcPr.remove(ex)
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        borders.append(el)
    ex = tcPr.find(qn("w:tcBorders"))
    if ex is not None:
        tcPr.remove(ex)
    tcPr.append(borders)


def hdr(cell, text, size=9):
    set_cell_bg(cell, HDR_FILL)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def dat(cell, text, alt=False, bold=False,
        align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=None):
    if alt:
        set_cell_bg(cell, ALT_FILL)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.alignment = align


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.clear()
    run = p.add_run(text)
    run.font.name  = FONT
    run.bold       = False
    run.italic     = True
    run.font.color.rgb = TITLE_RGB
    return p


def add_para(doc, text, style="Normal", size=None, color=None, bold=False):
    p = doc.add_paragraph(style=style)
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    return p


# ── data loader ────────────────────────────────────────────────────────────────

def _load_perf(path, drug_type):
    wb = openpyxl.load_workbook(path, read_only=True)
    ws = wb["ALL FACILITIES"]
    rows_raw = [
        r for r in ws.iter_rows(min_row=3, values_only=True)
        if r[2] and str(r[2]).strip() not in ("", "TOTAL", "GRAND TOTAL")
    ]
    wb.close()

    lga_d = defaultdict(lambda: dict(
        hfs=0, target=0, treated=0, records=0,
        dups=0, missing_hh=0, missing_child=0,
        age59=0, age0=0,
        absent=0, refused=0, inelig=0, referred=0, died=0, migrated=0,
        drug1=0, drug2=0,
    ))
    facilities = []

    for row in rows_raw:
        # col order: #, lga, fac, tgt, rec, treated, not_treated, drug1, drug2, cov, status,
        #            absent, refused, inelig, referred, died, migrated, redose,
        #            age59, age0, missing_hh, missing_child, dups, del_com, wards
        num, lga, fac, tgt, rec, treated, not_treated, drug1, drug2, cov, status = row[:11]
        (absent, refused, inelig, referred, died, migrated, redose,
         age59, age0, missing_hh, missing_child, dups, del_com) = row[11:24]

        def i(v): return int(v or 0)

        if not lga:
            continue

        L = lga_d[str(lga).strip()]
        L["hfs"]          += 1
        L["target"]       += i(tgt)
        L["treated"]      += i(treated)
        L["records"]      += i(rec)
        L["dups"]         += i(dups)
        L["missing_hh"]   += i(missing_hh)
        L["missing_child"]+= i(missing_child)
        L["age59"]        += i(age59)
        L["age0"]         += i(age0)
        L["absent"]       += i(absent)
        L["refused"]      += i(refused)
        L["inelig"]       += i(inelig)
        L["referred"]     += i(referred)
        L["died"]         += i(died)
        L["migrated"]     += i(migrated)
        L["drug1"]        += i(drug1)
        L["drug2"]        += i(drug2)

        facilities.append({
            "lga": str(lga).strip(), "fac": str(fac).strip(),
            "tgt": i(tgt), "rec": i(rec), "treated": i(treated),
            "cov": cov or "—", "status": str(status or "").strip(),
        })

    return lga_d, facilities


def _load_sync_summary(path):
    """Returns (lga_rows, time_stats) where time_stats = {label: (count, pct)}."""
    if not path or not os.path.exists(path):
        return [], {}
    wb   = openpyxl.load_workbook(path, read_only=True)
    ws   = wb["SUMMARY"]
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    wb.close()

    lga_rows   = []
    time_stats = {}
    for r in rows:
        if not r[0] and not r[1]:
            continue
        label = str(r[0] or "").strip()
        if "Synced by" in label:
            time_stats[label] = (r[1], r[2])   # (count, pct_str)
        elif r[1] and str(r[1]).strip().upper() not in ("GRAND TOTAL", "TOTAL"):
            lga_rows.append(r)
    return lga_rows, time_stats


def _load_facility_sync_rates(sync_path):
    """
    Read per-LGA tabs from the sync Excel and compute facility-level sync rates.
    Returns list of dicts sorted by sync_rate ascending (lowest first).
    """
    if not sync_path or not os.path.exists(sync_path):
        return []
    try:
        wb   = openpyxl.load_workbook(sync_path, read_only=True)
        skip = {"SUMMARY", "NEVER SYNCED", "LOW SYNCED"}
        fac_stats = {}   # facility -> {total, synced}

        for ws in wb.worksheets:
            if ws.title in skip:
                continue
            headers = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
            try:
                fac_col    = headers.index("Health Facility")
                lga_col    = headers.index("LGA")
                status_col = headers.index("Status")
            except ValueError:
                continue

            for row in ws.iter_rows(min_row=2, values_only=True):
                if not row or len(row) <= max(fac_col, lga_col, status_col):
                    continue
                fac    = str(row[fac_col] or "").strip()
                lga    = str(row[lga_col] or "").strip()
                status = str(row[status_col] or "").strip()
                if not fac:
                    continue
                key = (lga, fac)
                if key not in fac_stats:
                    fac_stats[key] = {"lga": lga, "fac": fac, "total": 0, "synced": 0}
                fac_stats[key]["total"] += 1
                if status != "NEVER SYNCED":
                    fac_stats[key]["synced"] += 1

        wb.close()
        results = []
        for v in fac_stats.values():
            rate = v["synced"] / v["total"] * 100 if v["total"] else 0
            results.append({**v, "rate": rate,
                             "rate_str": f"{rate:.1f}%",
                             "never": v["total"] - v["synced"]})
        return sorted(results, key=lambda x: x["rate"])   # lowest sync rate first
    except Exception as e:
        log.warning(f"[report] facility sync rate load failed (non-fatal): {e}")
        return []


def _load_all_days_perf(cfg):
    """
    Read each day's performance Excel and return a list of daily totals dicts.
    Used for the day-by-day comparison table and bar chart.
    """
    days = []
    for day_num in range(1, cfg["DAY"] + 1):
        path = os.path.join(cfg["out_dir"], f"performance_day{day_num}.xlsx")
        if not os.path.exists(path):
            continue
        try:
            wb = openpyxl.load_workbook(path, read_only=True)
            ws = wb["ALL FACILITIES"]
            rows_raw = [
                r for r in ws.iter_rows(min_row=3, values_only=True)
                if r[2] and str(r[2]).strip() not in ("", "TOTAL", "GRAND TOTAL")
            ]
            wb.close()
            records = sum(int(r[4] or 0) for r in rows_raw)
            treated = sum(int(r[5] or 0) for r in rows_raw)
            target  = sum(int(r[3] or 0) for r in rows_raw)
            cov_pct = treated / target * 100 if target else 0
            date_label = (cfg["campaign_start"] + timedelta(days=day_num - 1)).strftime("%d %b")
            days.append({
                "day":      day_num,
                "date":     date_label,
                "records":  records,
                "treated":  treated,
                "target":   target,
                "cov_pct":  cov_pct,
                "coverage": f"{cov_pct:.1f}%" if target else "N/A",
            })
        except Exception as e:
            log.warning(f"Could not read performance_day{day_num}.xlsx: {e}")
    return days


def _generate_progress_chart(days_data, cfg):
    """Single-panel daily coverage % bar chart."""
    if not days_data:
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        labels   = [f"Day {d['day']}\n{d['date']}" for d in days_data]
        coverage = [d["cov_pct"] for d in days_data]

        def bar_color(c):
            if c >= 95: return "#1A7A1A"
            if c >= 70: return "#E06000"
            return "#CC0000"

        fig, ax = plt.subplots(figsize=(9, 4))
        fig.patch.set_facecolor("#F9F9F9")

        colors = [bar_color(c) for c in coverage]
        bars   = ax.bar(labels, coverage, color=colors, width=0.5, zorder=3)
        ax.set_title(f"{cfg['state_name']} — Daily Coverage %",
                     fontsize=11, fontweight="bold", pad=10)
        ax.set_ylabel("Coverage (%)", fontsize=9)
        ax.set_ylim(0, 115)
        ax.axhline(95, color="#1A7A1A", linestyle="--", linewidth=1, alpha=0.6, label="HIGH (95%)")
        ax.axhline(70, color="#E06000", linestyle="--", linewidth=1, alpha=0.6, label="MODERATE (70%)")
        ax.legend(fontsize=8)
        ax.grid(axis="y", linestyle="--", alpha=0.4, zorder=0)
        ax.set_axisbelow(True)
        for bar, val in zip(bars, coverage):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1.5,
                    f"{val:.1f}%", ha="center", va="bottom", fontsize=9, fontweight="bold")

        plt.tight_layout(pad=2.0)
        chart_path = os.path.join(cfg["out_dir"], f"progress_chart_day{cfg['DAY']}.png")
        plt.savefig(chart_path, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
        plt.close()
        log.info(f"[report] progress chart saved -> {chart_path}")
        return chart_path
    except Exception as e:
        log.warning(f"[report] chart generation failed (non-fatal): {e}")
        return None


def _grand_totals(lga_d):
    g = defaultdict(int)
    for L in lga_d.values():
        for k, v in L.items():
            g[k] += v
    return g


# ── previous report reader ─────────────────────────────────────────────────────

def _read_previous_report(cfg):
    """
    Find the most recent earlier report .docx to use as context for Claude.

    Priority:
      1. Latest report from the same campaign day (earlier HHMM than now)
      2. If none, latest report from the previous campaign day (Day N-1)

    Returns the full paragraph text, or empty string if nothing found.
    """
    import glob as _glob
    from datetime import datetime

    state  = str(cfg.get("state_name", "")).strip().replace(" ", "_")
    day    = cfg["DAY"]
    folder = cfg["out_dir"]
    now_hm = datetime.now().strftime("%H%M")

    # 1. Same-day reports (earlier HHMM)
    pattern   = os.path.join(folder, f"{state}_Day{day}_Report_*.docx")
    same_day  = []
    for fp in sorted(_glob.glob(pattern)):
        tag = os.path.basename(fp).replace(f"{state}_Day{day}_Report_", "").replace(".docx", "")
        if tag.isdigit() and tag < now_hm:
            same_day.append(fp)

    if same_day:
        prev_path = same_day[-1]
    else:
        # 2. Fall back to last report of previous day
        prev_day     = day - 1
        prev_pattern = os.path.join(folder, f"{state}_Day{prev_day}_Report_*.docx")
        prev_day_files = sorted(_glob.glob(prev_pattern))
        if not prev_day_files:
            return ""
        prev_path = prev_day_files[-1]   # latest HHMM of previous day

    log.info(f"[report] reading previous report for context: {prev_path}")
    try:
        prev_doc = Document(prev_path)
        text = "\n".join(p.text for p in prev_doc.paragraphs if p.text.strip())
        return text
    except Exception as e:
        log.warning(f"[report] could not read previous report (non-fatal): {e}")
        return ""


# ── Claude API ─────────────────────────────────────────────────────────────────

def _claude(prompt, max_tokens=400):
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        log.warning("ANTHROPIC_API_KEY not set — returning placeholder text")
        return "[Narrative not generated — ANTHROPIC_API_KEY missing]"
    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}],
        )
        return resp.content[0].text.strip()
    except Exception as e:
        log.warning(f"Claude API call failed (non-fatal): {e}")
        return "[Narrative not generated — Claude API error]"


def _conclusion_prompt(cfg, g, cov_pct, lga_d, sync_rows, sync_time_stats, prev_report):
    best_lga  = max(lga_d, key=lambda l: lga_d[l]["treated"] / lga_d[l]["target"] * 100
                    if lga_d[l]["target"] else 0, default="N/A")
    worst_lga = min(lga_d, key=lambda l: lga_d[l]["treated"] / lga_d[l]["target"] * 100
                    if lga_d[l]["target"] else 100, default="N/A")
    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if r[2])
    synced_cdds = sum(int(r[3] or 0) for r in sync_rows if r[3])
    sync_pct    = f"{synced_cdds/total_cdds*100:.1f}%" if total_cdds else "N/A"
    by_17_str   = ""
    if sync_time_stats:
        for label, (count, pct) in sync_time_stats.items():
            by_17_str = f"\n- Synced by 17:00: {count:,} ({pct})"

    prev_block = ""
    if prev_report:
        prev_block = (f"\n--- PREVIOUS REPORT (compare — show progress) ---\n"
                      f"{prev_report}\n--- END ---\n")

    return f"""You are writing the conclusion of a formal health campaign daily report.
Campaign: {cfg['campaign_name']} — {cfg['state_name']}
Day {cfg['DAY']} of {cfg['campaign_days']} ({cfg['DATE_LABEL']})
Drug: {cfg['drug_type']}
{prev_block}
Statistics:
- Daily target: {g['target']:,} | Records: {g['records']:,} | Treated: {g['treated']:,} | Coverage: {cov_pct}
- Best LGA: {best_lga} | Weakest LGA: {worst_lga}
- Duplicates: {g['dups']:,} | Missing child: {g['missing_child']:,}
- FLW sync: {synced_cdds:,}/{total_cdds:,} ({sync_pct}){by_17_str}

Write a conclusion paragraph of EXACTLY 5 sentences:
1. Overall coverage vs daily target{"— compare to previous report" if prev_report else ""}.
2. Best-performing LGA with specific numbers.
3. Weakest LGA and operational implication.
4. Main data quality concern and recommended action.
5. FLW sync status and outlook for the next extract.

Rules: formal tone, plain text only, no bullet points, no headings, no emojis."""


def _issues_prompt(cfg, g, cov_pct, lga_d, facilities, sync_rows, sync_time_stats, prev_report):
    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if r[2])
    synced_cdds = sum(int(r[3] or 0) for r in sync_rows if r[3])
    never       = sum(int(r[6] or 0) for r in sync_rows if r[6])
    sync_pct    = f"{synced_cdds/total_cdds*100:.1f}%" if total_cdds else "N/A"

    worst_sync = sorted(
        [(str(r[1]), int(r[6] or 0), int(r[2] or 0))
         for r in sync_rows if r[1] and int(r[2] or 0) > 0],
        key=lambda x: -x[1],
    )
    worst_sync_str = (
        f"{worst_sync[0][0]} worst: {worst_sync[0][1]} of {worst_sync[0][2]} never synced"
        if worst_sync else "N/A"
    )
    by_17_str = ""
    if sync_time_stats:
        for label, (count, pct) in sync_time_stats.items():
            by_17_str = f"  Synced by 17:00: {count:,} ({pct})"

    low_act = [f for f in facilities if f["rec"] < 10]
    is_last = cfg["DAY"] == cfg["campaign_days"]
    prev_block = (f"\n--- PREVIOUS REPORT (compare — mark issues RESOLVED if fixed) ---\n"
                  f"{prev_report}\n--- END ---\n") if prev_report else ""

    return f"""You are generating a structured issues log for a health campaign operations report.

Campaign: {cfg['campaign_name']} — {cfg['state_name']}
Day {cfg['DAY']} of {cfg['campaign_days']} ({'FINAL day' if is_last else 'ongoing'}) — {cfg['DATE_LABEL']}
{prev_block}
KEY METRICS:
- Coverage: {cov_pct} ({g['treated']:,} treated / {g['target']:,} target)
- CDDs: {total_cdds:,} total | {synced_cdds:,} synced ({sync_pct}) | {never:,} never synced
{by_17_str}
- {worst_sync_str}
- Low activity facilities (<10 records): {len(low_act)}
- Duplicates: {g['dups']:,} | Missing child: {g['missing_child']:,}

Generate 3-5 issues. Return ONLY a JSON array — no markdown, no explanation:
[
  {{
    "observation": "Concise fact with specific numbers and LGA names",
    "status": "ACTIVE",
    "priority": "High",
    "notes": "Specific action for field supervisor or coordinator",
    "data_type": "perf"
  }}
]

Rules:
- status: "ACTIVE" unless previous report shows issue resolved → "RESOLVED"
- priority: "High" (blocks target), "Moderate" (needs attention), "Low" (monitor)
- data_type: "sync" for CDD/sync issues, "perf" for coverage/facility/data-quality issues
- observation: include specific numbers, LGA names, percentages
- notes: actionable, field-level instruction"""


def _claude_issues(cfg, g, cov_pct, lga_d, facilities, sync_rows, sync_time_stats, prev_report):
    import json
    prompt = _issues_prompt(cfg, g, cov_pct, lga_d, facilities, sync_rows, sync_time_stats, prev_report)
    raw = _claude(prompt, max_tokens=900)
    try:
        text = raw.strip()
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:])
            if text.rstrip().endswith("```"):
                text = text.rstrip()[:-3]
        return json.loads(text)
    except Exception as e:
        log.warning(f"[report] issues JSON parse failed: {e} — using raw text as single issue")
        return [{
            "observation": raw[:300] if raw else "No issues generated.",
            "status": "ACTIVE",
            "priority": "High",
            "notes": "Review manually.",
            "data_type": "perf",
        }]


def _slack_prompt(cfg, g, cov_pct, docx_name, sync_rows, sync_time_stats, prev_report):
    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if r[2])
    synced_cdds = sum(int(r[3] or 0) for r in sync_rows if r[3])
    sync_pct    = f"{synced_cdds/total_cdds*100:.1f}%" if total_cdds else "N/A"
    by_17_str = ""
    if sync_time_stats:
        for label, (count, pct) in sync_time_stats.items():
            by_17_str = f"  Synced by 17:00: {count:,} ({pct})"

    prev_block = (f"\n--- PREVIOUS REPORT ---\n{prev_report}\n--- END ---\n") if prev_report else ""

    return f"""You are writing a Slack notification for a health campaign operations team.

Campaign: {cfg['campaign_name']} — {cfg['state_name']}
Day {cfg['DAY']} of {cfg['campaign_days']}, {cfg['DATE_LABEL']}
{prev_block}
Stats: Target={g['target']:,} | Records={g['records']:,} | Treated={g['treated']:,} | Coverage={cov_pct}
Data quality: Dups={g['dups']:,} | Missing child={g['missing_child']:,}
Sync: {synced_cdds:,}/{total_cdds:,} ({sync_pct}){by_17_str}

Write the Slack message in EXACTLY this format:
Hi Team,

[Line 1: coverage result and treated count — one sentence{"— compare to previous report" if prev_report else ""}]
[Line 2: FLW sync status including synced-by-17:00 figure and most urgent action — one sentence]

Rules: plain text only, no emojis, no bullet points, no asterisks, no headings. The link and footer are appended automatically — do not include them."""


# ── document builder ───────────────────────────────────────────────────────────

def _cov_str(treated, target):
    if not target:
        return "N/A"
    return f"{treated/target*100:.1f}%"


def _two_col_table(doc, rows_data, col_widths=(5, 9)):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    for ri, (param, val) in enumerate(rows_data):
        hdr(table.cell(ri, 0), param, size=9)
        dat(table.cell(ri, 1), val,   alt=(ri % 2 == 1), bold=False,
            align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
    table.columns[0].width = Cm(col_widths[0])
    table.columns[1].width = Cm(col_widths[1])


def _perf_table(doc, lga_d, lga_summary=True):
    header = ["LGA", "HFs", "Target", "Treated", "Coverage %", "Status"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, (lga, L) in enumerate(sorted(lga_d.items()), 1):
        raw_pct = L["treated"] / L["target"] * 100 if L["target"] else None
        cov     = f"{raw_pct:.1f}%" if raw_pct is not None else "N/A"
        stat    = _cov_band(raw_pct) if raw_pct is not None else "NO TARGET"
        vals    = [lga, L["hfs"], f"{L['target']:,}", f"{L['treated']:,}", cov, stat]
        row     = table.add_row()
        alt     = ri % 2 == 1
        for ci, val in enumerate(vals):
            if alt:
                set_cell_bg(row.cells[ci], ALT_FILL)
            set_cell_borders(row.cells[ci])
            p = row.cells[ci].paragraphs[0]
            p.clear()
            run = p.add_run(str(val))
            run.font.name = FONT; run.font.size = Pt(9)
            if ci in (4, 5):
                run.bold = True
                color = STATUS_COLOR.get(stat)
                if color: run.font.color.rgb = color
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER


def _dq_table(doc, lga_d):
    header = ["LGA", "Duplicates", "Missing HH", "Missing Child", "Age=0", "Age>59"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, (lga, L) in enumerate(sorted(lga_d.items()), 1):
        vals = [lga, L["dups"], L["missing_hh"], L["missing_child"], L["age0"], L["age59"]]
        row  = table.add_row()
        alt  = ri % 2 == 1
        for ci, val in enumerate(vals):
            dat(row.cells[ci], val, alt=alt)


def _fac_perf_table(doc, facilities, perf_link=""):
    """
    LOW coverage (<70%) and Low Activity (<10 records) facilities.
    Pop. Coverage = Records / Daily Target.
    Treatment Coverage = Treated / Records.
    Cells colour-coded: >=95% green, 70-95% amber, <70% red.
    """
    low_facs = [
        f for f in facilities
        if f["status"] in ("LOW", "LOW ACTIVITY") and f["tgt"] > 0
    ]
    low_facs = sorted(low_facs, key=lambda f: f["rec"] / f["tgt"] if f["tgt"] else 0)

    total_facs = len(facilities)
    shown      = min(len(low_facs), 20)
    note_text  = f"Showing {shown} LOW / Low Activity facilities of {total_facs} total. "
    note_p     = add_para(doc, note_text, size=9, color=GREY_RGB)
    if perf_link:
        _add_hyperlink(note_p, "Full list in performance Excel ↗", perf_link)
    else:
        run = note_p.add_run("Full list in performance Excel.")
        run.font.name = FONT; run.font.size = Pt(9); run.font.color.rgb = GREY_RGB

    if not low_facs:
        add_para(doc, "No LOW or Low Activity facilities on this day.")
        return

    header = ["#", "District", "Health Facility", "Daily Target", "Records",
              "Treated", "Not Treated", "Pop. Coverage", "Treatment Coverage"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)

    for ri, f in enumerate(low_facs[:20], 1):
        pop_pct  = f["rec"]     / f["tgt"] * 100 if f["tgt"] else None
        trt_pct  = f["treated"] / f["rec"] * 100 if f["rec"] else None
        pop_str  = f"{pop_pct:.1f}%"  if pop_pct  is not None else "N/A"
        trt_str  = f"{trt_pct:.1f}%"  if trt_pct  is not None else "N/A"
        not_trt  = f["rec"] - f["treated"]

        pop_band = _cov_band(pop_pct) if pop_pct is not None else "NO TARGET"
        trt_band = _cov_band(trt_pct) if trt_pct is not None else "NO TARGET"

        row = table.add_row()
        alt = ri % 2 == 1
        vals = [ri, f["lga"], f["fac"], f"{f['tgt']:,}", f"{f['rec']:,}",
                f"{f['treated']:,}", f"{not_trt:,}", pop_str, trt_str]

        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            if alt:
                set_cell_bg(cell, ALT_FILL)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.clear()
            r = p.add_run(str(val))
            r.font.name = FONT; r.font.size = Pt(9)
            # coloured text only on the two coverage columns
            if ci == 7:
                r.bold = True
                if STATUS_COLOR.get(pop_band): r.font.color.rgb = STATUS_COLOR[pop_band]
            if ci == 8:
                r.bold = True
                if STATUS_COLOR.get(trt_band): r.font.color.rgb = STATUS_COLOR[trt_band]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci in (1, 2) else WD_ALIGN_PARAGRAPH.CENTER


def _low_activity_table(doc, facilities):
    low_act = [f for f in facilities if f["rec"] < 10]
    if not low_act:
        add_para(doc, "No facilities with fewer than 10 records on this day.")
        return
    header = ["Health Facility", "LGA", "Target", "Records", "Treated", "Status"]
    table  = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for ci, h in enumerate(header):
        hdr(table.cell(0, ci), h)
    for ri, f in enumerate(sorted(low_act, key=lambda x: x["rec"]), 1):
        row = table.add_row()
        alt = ri % 2 == 1
        dat(row.cells[0], f["fac"],     alt=alt, align=WD_ALIGN_PARAGRAPH.LEFT)
        dat(row.cells[1], f["lga"],     alt=alt)
        dat(row.cells[2], f"{f['tgt']:,}", alt=alt)
        dat(row.cells[3], f"{f['rec']:,}", alt=alt)
        dat(row.cells[4], f"{f['treated']:,}", alt=alt)
        dat(row.cells[5], f["status"], alt=alt,
            color=STATUS_COLOR.get(f["status"]))


def _non_admin_table(doc, g):
    reasons = [
        ("Absent",      g["absent"]),
        ("Refused",     g["refused"]),
        ("Ineligible",  g["inelig"]),
        ("Referred",    g["referred"]),
        ("Died",        g["died"]),
        ("Migrated",    g["migrated"]),
    ]
    reasons = [(r, c) for r, c in reasons if c > 0]
    if not reasons:
        add_para(doc, "No non-administration events recorded.")
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr(table.cell(0, 0), "Reason")
    hdr(table.cell(0, 1), "Count")
    for ri, (reason, count) in enumerate(sorted(reasons, key=lambda x: -x[1]), 1):
        row = table.add_row()
        dat(row.cells[0], reason, alt=ri % 2 == 1, align=WD_ALIGN_PARAGRAPH.LEFT)
        dat(row.cells[1], f"{count:,}", alt=ri % 2 == 1)


def _dq_summary_table(doc, g):
    total = g["treated"] or 1
    metrics = [
        ("Duplicate Records",    g["dups"]),
        ("Missing HH Name",      g["missing_hh"]),
        ("Missing Child Name",   g["missing_child"]),
        ("Age = 0",              g["age0"]),
        ("Age > 59 months",      g["age59"]),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr(table.cell(0, 0), "Metric")
    hdr(table.cell(0, 1), "Count")
    hdr(table.cell(0, 2), "% of Treated")
    for ri, (metric, count) in enumerate(metrics, 1):
        row = table.add_row()
        pct = f"{count/total*100:.2f}%"
        dat(row.cells[0], metric, alt=ri % 2 == 1, align=WD_ALIGN_PARAGRAPH.LEFT)
        dat(row.cells[1], f"{count:,}", alt=ri % 2 == 1)
        dat(row.cells[2], pct, alt=ri % 2 == 1)


def _sync_table(doc, sync_rows, cfg, sync_time_stats=None):
    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if r[2])
    synced_cdds = sum(int(r[3] or 0) for r in sync_rows if r[3])
    never       = sum(int(r[6] or 0) for r in sync_rows if r[6])
    sync_pct    = f"{synced_cdds/total_cdds*100:.1f}%" if total_cdds else "N/A"

    add_heading(doc, "5.1  FLW Sync Summary", 5)
    summary_rows = [
        ("Total CDDs Registered",   f"{total_cdds:,}"),
        ("CDDs Synced (total day)",  f"{synced_cdds:,}  ({sync_pct})"),
        ("Never Synced",             f"{never:,}"),
    ]
    # Time-based breakdown
    if sync_time_stats:
        for label, (count, pct) in sync_time_stats.items():
            hour = label.replace("Synced by ", "").replace(" today (UTC)", "")
            if count is not None:
                summary_rows.append(
                    (f"Synced by {hour}",
                     f"{count:,}  ({pct})  — early sync indicator")
                )
    summary_rows += [
        ("Report Date",  cfg["DATE_LABEL"]),
    ]
    _two_col_table(doc, summary_rows, col_widths=(5, 6))
    doc.add_paragraph()

    if sync_rows:
        add_heading(doc, "5.1b  Sync Status by LGA", 5)
        cols = ["#", "LGA", "Total CDDs", "HIGH", "MODERATE", "LOW", "NEVER SYNCED", "% Never Synced"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for ci, h in enumerate(cols):
            hdr(table.cell(0, ci), h)
        for ri, row in enumerate(sync_rows, 1):
            tr = table.add_row()
            alt = ri % 2 == 1
            for ci in range(len(cols)):
                val = ri if ci == 0 else (row[ci - 1] if ci - 1 < len(row) else "")
                dat(tr.cells[ci], val, alt=alt)


# ── public entry point ─────────────────────────────────────────────────────────

def run(cfg):
    log.info(f"[report] {cfg['state_name']} Day {cfg['DAY']} ...")

    perf_path = cfg["perf_xlsx"]
    sync_path = cfg["sync_xlsx"]

    if not os.path.exists(perf_path):
        raise FileNotFoundError(f"Performance Excel not found: {perf_path}")

    drug_type  = cfg["drug_type"]
    d1_label   = "SPAQ1 (12-59m)" if drug_type == "SPAQ" else "AZM 12-59m"
    d2_label   = "SPAQ2 (3-11m)"  if drug_type == "SPAQ" else "AZM 1-11m"

    lga_d, facilities          = _load_perf(perf_path, drug_type)
    sync_rows, sync_time_stats = _load_sync_summary(sync_path)
    g                          = _grand_totals(lga_d)
    cov_pct           = _cov_str(g["treated"], g["target"])
    hfs_active        = len({f["lga"] for f in facilities})

    log.info(f"  {len(facilities)} facilities, {len(lga_d)} LGAs, coverage {cov_pct}")

    # Load day-by-day totals for cumulative stats + chart
    days_data   = _load_all_days_perf(cfg)
    cum_records = sum(d["records"] for d in days_data)
    cum_treated = sum(d["treated"] for d in days_data)
    cum_target  = sum(d["target"]  for d in days_data)
    cum_cov     = f"{cum_treated/cum_target*100:.1f}%" if cum_target else "N/A"
    chart_path  = _generate_progress_chart(days_data, cfg)

    # Read previous report for trend context
    prev_report = _read_previous_report(cfg)
    if prev_report:
        log.info(f"  previous report loaded ({len(prev_report)} chars)")
    else:
        log.info("  no previous report — first extract today")

    # Upload raw Excel files to Drive so rows in the issues table can link to them
    perf_link = ""
    sync_link = ""
    try:
        import notify as _notify
        _now_hm    = datetime.now().strftime("%H:%M")
        perf_title = f"{cfg['state_name']} Day {cfg['DAY']} Performance Data — {cfg['DATE_LABEL']} {_now_hm}"
        sync_title = f"{cfg['state_name']} Day {cfg['DAY']} CDD Sync Data — {cfg['DATE_LABEL']} {_now_hm}"
        log.info("  uploading performance Excel to Drive ...")
        perf_link = _notify.upload_file(perf_path, perf_title)
        if sync_path and os.path.exists(sync_path):
            log.info("  uploading CDD sync Excel to Drive ...")
            sync_link = _notify.upload_file(sync_path, sync_title)
    except Exception as e:
        log.warning(f"  Drive upload of Excels failed (non-fatal): {e}")

    # Claude calls
    log.info("  calling Claude for issues log ...")
    issues_data = _claude_issues(cfg, g, cov_pct, lga_d, facilities,
                                 sync_rows, sync_time_stats, prev_report)
    log.info(f"  {len(issues_data)} issues generated")
    log.info("  calling Claude for conclusion ...")
    conclusion = _claude(_conclusion_prompt(cfg, g, cov_pct, lga_d, sync_rows, sync_time_stats, prev_report),
                         max_tokens=500)
    log.info("  calling Claude for Slack text ...")
    slack_text = _claude(
        _slack_prompt(cfg, g, cov_pct, os.path.basename(cfg["docx_path"]),
                      sync_rows, sync_time_stats, prev_report),
        max_tokens=300,
    )

    # ── Build document ────────────────────────────────────────────────────────
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Title block — centred, navy, Times New Roman (matches tg_smc reference)
    title_p = doc.add_paragraph(style="Normal")
    title_p.clear()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"{cfg['state_name']}  —  {cfg['campaign_name']}")
    run.font.name = FONT; run.font.size = Pt(26)
    run.bold = True; run.font.color.rgb = TITLE_RGB

    sub = doc.add_paragraph(style="Normal")
    sub.clear()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"{cfg['START_LABEL']} to {cfg['END_LABEL']}  —  Day {cfg['DAY']} of {cfg['campaign_days']}")
    r2.font.name = FONT; r2.font.size = Pt(14)
    r2.bold = True; r2.font.color.rgb = TITLE_RGB

    summary_line = (
        f"Day {cfg['DAY']}  —  {cfg['DATE_LABEL']}  |  "
        f"Extract: {cfg['DATE_LABEL']}, {datetime.now().strftime('%H:%M')}  |  "
        f"{hfs_active} LGAs  |  Coverage: {cov_pct}"
    )
    grey_p = add_para(doc, summary_line, size=8, color=GREY_RGB)
    grey_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Sync totals needed for overview table
    total_cdds  = sum(int(r[2] or 0) for r in sync_rows if r[2])
    synced_cdds = sum(int(r[3] or 0) for r in sync_rows if r[2])
    never_cdds  = sum(int(r[6] or 0) for r in sync_rows if r[6])
    sync_pct_ov = f"{synced_cdds/total_cdds*100:.1f}%" if total_cdds else "N/A"
    by_17_val   = "—"
    if sync_time_stats:
        for label, (count, pct) in sync_time_stats.items():
            by_17_val = f"{count:,}  ({pct})" if count is not None else "—"
    not_admin = g['absent'] + g['refused'] + g['inelig'] + g['referred'] + g['died'] + g['migrated']

    # Section 1
    add_heading(doc, f"1.  Day {cfg['DAY']} Operational Overview", 4)
    overview_rows = [
        ("State / Country",               cfg['state_name']),
        ("Activity",                      f"{cfg['drug_type']} Distribution  —  Day {cfg['DAY']} of {cfg['campaign_days']}"),
        ("Date",                          cfg['DATE_LABEL']),
        ("Data Extract Timestamp",        f"{cfg['DATE_LABEL']}, {datetime.now().strftime('%H:%M')}"),
        ("Campaign Dates",                f"{cfg['START_LABEL']} to {cfg['END_LABEL']}"),
        ("LGAs / Districts Covered",      f"{hfs_active} of {cfg['hfs_total']}"),
        # ── Coverage ───────────────────────────────────────────────────────
        ("Daily Population Target",       f"{g['target']:,}"),
        ("Total Records Submitted",       f"{g['records']:,}"),
        ("Children Treated",              f"{g['treated']:,}"),
        ("Coverage vs Daily Target",      cov_pct),
        # ── Cumulative ─────────────────────────────────────────────────────
        (f"Cumulative Records (Days 1–{cfg['DAY']})",  f"{cum_records:,}"),
        (f"Cumulative Treated (Days 1–{cfg['DAY']})",  f"{cum_treated:,}"),
        (f"Cumulative Coverage (Days 1–{cfg['DAY']})", cum_cov),
        # ── Drug split ─────────────────────────────────────────────────────
        (d1_label,                        f"{g['drug1']:,}"),
        (d2_label,                        f"{g['drug2']:,}"),
        ("Not Administered",              f"{not_admin:,}"),
        # ── Sync ───────────────────────────────────────────────────────────
        ("CDDs Registered",               f"{total_cdds:,}"),
        ("CDDs Synced",                   f"{synced_cdds:,}  ({sync_pct_ov})"),
        ("CDDs Synced by 17:00",          by_17_val),
    ]
    _two_col_table(doc, overview_rows)
    doc.add_paragraph()

    # Coverage chart — directly under Section 1 overview table
    if chart_path and os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(5.5))
        doc.add_paragraph()

    # Section 2 — Issues log (Claude-generated, with Drive links per row)
    add_heading(doc, f"2.  Program Issues and Resolutions — Day {cfg['DAY']} Log", 4)
    add_para(doc,
             f"Day {cfg['DAY']} of {cfg['campaign_days']}. Issues identified during the extract and field reports.",
             size=9, color=GREY_RGB)

    issues_header = ["#", "Issue / Observation", "Status", "Priority", "Notes", "Data"]
    tbl2 = doc.add_table(rows=1, cols=len(issues_header))
    tbl2.style = "Table Grid"
    for ci, h in enumerate(issues_header):
        hdr(tbl2.cell(0, ci), h)

    _ACTIVE_RED   = RGBColor(0xCC, 0x00, 0x00)
    _RESOLVED_GRN = RGBColor(0x1A, 0x7A, 0x1A)
    _PRI_HIGH     = RGBColor(0xCC, 0x00, 0x00)
    _PRI_MOD      = RGBColor(0xE0, 0x60, 0x00)

    for ri, issue in enumerate(issues_data, 1):
        row2 = tbl2.add_row()
        alt  = ri % 2 == 1
        status   = str(issue.get("status",   "ACTIVE")).upper()
        priority = str(issue.get("priority", "High"))
        dtype    = str(issue.get("data_type","perf")).lower()

        dat(row2.cells[0], ri, alt=alt)
        dat(row2.cells[1], issue.get("observation", ""), alt=alt,
            align=WD_ALIGN_PARAGRAPH.LEFT, size=9)
        dat(row2.cells[2], status, alt=alt,
            color=_ACTIVE_RED if status == "ACTIVE" else _RESOLVED_GRN)
        dat(row2.cells[3], priority, alt=alt,
            color=_PRI_HIGH if priority == "High" else _PRI_MOD if priority == "Moderate" else None)
        dat(row2.cells[4], issue.get("notes", ""), alt=alt,
            align=WD_ALIGN_PARAGRAPH.LEFT, size=9)

        # Data column — hyperlink to relevant Drive file
        data_cell = row2.cells[5]
        if alt:
            set_cell_bg(data_cell, ALT_FILL)
        set_cell_borders(data_cell)
        dp = data_cell.paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        link  = sync_link  if dtype == "sync"  else perf_link
        label = "CDD Sync Data" if dtype == "sync" else "Performance Data"
        fname = (os.path.basename(sync_path)  if dtype == "sync"  else os.path.basename(perf_path)) or label

        if link:
            _add_hyperlink(dp, label + " ↗", link)
        else:
            run_d = dp.add_run(fname)
            run_d.font.name = FONT; run_d.font.size = Pt(8)
            run_d.font.color.rgb = GREY_RGB

    # column widths: #, Observation, Status, Priority, Notes, Data
    for ci, w in enumerate([Cm(0.8), Cm(5.5), Cm(1.5), Cm(1.6), Cm(4.5), Cm(2.5)]):
        for cell in tbl2.columns[ci].cells:
            cell.width = w
    doc.add_paragraph()

    # Section 3
    add_heading(doc, "3.  Distribution Data Analysis", 4)

    add_heading(doc, "3.1  Performance by LGA", 5)
    _perf_table(doc, lga_d)
    doc.add_paragraph()

    add_heading(doc, "3.2  Data Quality by LGA", 5)
    _dq_table(doc, lga_d)
    doc.add_paragraph()

    add_heading(doc, "3.3  Facility Performance Analysis", 5)
    add_para(doc, "LOW coverage (<70%) and Low Activity (<10 records) facilities. Sorted by Population Coverage ascending.", size=9, bold=True)
    _fac_perf_table(doc, facilities, perf_link=perf_link)
    doc.add_paragraph()

    add_heading(doc, "3.4  Non-Administration Analysis", 5)
    _non_admin_table(doc, g)
    doc.add_paragraph()

    add_heading(doc, "3.5  Data Quality Summary", 5)
    _dq_summary_table(doc, g)
    doc.add_paragraph()

    # Section 4 — Campaign Progress (day-by-day)
    add_heading(doc, f"4.  Campaign Progress  —  Days 1 to {cfg['DAY']}", 4)
    add_para(doc, "Coverage = Treated / Daily Target. Cumulative figures include all days to date.", size=9, color=GREY_RGB)

    if days_data:
        cols = ["Day", "Date", "Daily Target", "Records", "Treated", "Coverage"]
        tbl4 = doc.add_table(rows=1, cols=len(cols))
        tbl4.style = "Table Grid"
        for ci, h in enumerate(cols):
            hdr(tbl4.cell(0, ci), h)
        for ri, d in enumerate(days_data):
            row4 = tbl4.add_row()
            alt  = ri % 2 == 1
            cov_color = STATUS_COLOR.get(
                _cov_band(d["cov_pct"]) if d["target"] else "NO TARGET"
            )
            vals = [f"Day {d['day']}", d["date"], f"{d['target']:,}",
                    f"{d['records']:,}", f"{d['treated']:,}", d["coverage"]]
            for ci, val in enumerate(vals):
                dat(row4.cells[ci], val, alt=alt,
                    color=cov_color if ci == 5 else None)
        # Cumulative totals row
        tot_row = tbl4.add_row()
        for ci, val in enumerate(["TOTAL", "", f"{cum_target:,}", f"{cum_records:,}", f"{cum_treated:,}", cum_cov]):
            dat(tot_row.cells[ci], val, bold=True)

    doc.add_paragraph()

    # Section 5 — Sync
    add_heading(doc, "5.  Health Facility Data Synchronisation Status", 4)
    _sync_table(doc, sync_rows, cfg, sync_time_stats)

    # 5.2 Top 10 facilities with lowest CDD sync rate
    fac_sync = _load_facility_sync_rates(sync_path)
    if fac_sync:
        doc.add_paragraph()
        add_heading(doc, "5.2  Top 10 Facilities with Lowest CDD Sync Rate", 5)
        note_p = add_para(doc, "Facilities requiring immediate follow-up from supervisors. ", size=9, color=GREY_RGB)
        if sync_link:
            _add_hyperlink(note_p, "Full CDD sync data ↗", sync_link)
        low10  = fac_sync[:10]
        cols6  = ["#", "LGA", "Health Facility", "Total CDDs", "Synced", "Never Synced", "Sync Rate"]
        tbl6   = doc.add_table(rows=1, cols=len(cols6))
        tbl6.style = "Table Grid"
        for ci, h in enumerate(cols6):
            hdr(tbl6.cell(0, ci), h)
        for ri, f in enumerate(low10, 1):
            row6 = tbl6.add_row()
            alt  = ri % 2 == 1
            vals = [ri, f["lga"], f["fac"], f["total"], f["synced"], f["never"], f["rate_str"]]
            for ci, val in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if ci == 2 else WD_ALIGN_PARAGRAPH.CENTER
                dat(row6.cells[ci], val, alt=alt, align=align)

    # Section 6 — Conclusion (last section)
    doc.add_paragraph()
    add_heading(doc, "6.  Conclusion", 4)
    add_para(doc, conclusion, size=10)

    out = cfg["docx_path"]
    doc.save(out)
    log.info(f"[report] saved -> {out}")
    return out, slack_text
